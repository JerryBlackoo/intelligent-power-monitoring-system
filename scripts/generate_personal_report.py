from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
PROJECT = ROOT / "intelligent-power-monitoring-system"
TODAY = ROOT / "20260624"
OUT_PATH = TODAY / "个人实践报告_系统架构与前后端平台实现.docx"
SCREENSHOT = PROJECT / "static" / "reports" / "dashboard-smoke.png"


def set_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size: int = 11, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13, bold=True)


def add_para(doc: Document, text: str, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    if first_line:
        p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, size=11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r, size=11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = head
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_paragraph_font(paragraph, size=10)
    doc.add_paragraph()


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    r.font.size = Pt(9)


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def build_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    add_page_number(section)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("电力智能巡检系统个人实践报告")
    set_font(tr, size=20, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("系统总体架构设计与前后端平台实现")
    set_font(sr, size=14, bold=True)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("负责内容：系统总体架构设计、app 后端服务、static 前端工作台")
    set_font(mr, size=11)

    add_heading(doc, "摘要", 1)
    add_para(
        doc,
        "本实践项目面向电力设备巡检场景，设计并实现了基于边缘视觉推理与云端协同管理的电力智能巡检系统。系统采用“边缘节点 Atlas 200I DK / 摄像头 - FastAPI 云端后端 - Web 工作台”的三层架构，支持边端图片与推理结果上报、实时监控展示、告警处理、巡检记录管理、报告上传、Agent 辅助分析以及管理员后台概览等功能。本人主要负责系统总体架构设计，为组员提供可分工协作的接口、数据表和页面框架；同时负责 app 目录中的后端接口、数据持久化与业务服务，以及 static 目录中的前端工作台、角色页面、实时刷新和推理结果可视化。通过接口联调和页面测试，系统已能够保存边端图片、展示最新推理画面、绘制检测框、生成告警和巡检记录，并支持监控员、巡检员、管理员三类角色协作完成巡检闭环。",
    )
    add_para(doc, "关键词：电力巡检；边缘推理；FastAPI；YOLO；实时监控；告警处理；前后端协同", first_line=False)

    add_heading(doc, "1. 引言", 1)
    add_heading(doc, "1.1 项目背景与研究现状", 2)
    add_para(
        doc,
        "电力设备运行环境复杂，传统人工巡检存在效率低、记录滞后、异常发现依赖经验等问题。随着计算机视觉、深度学习和边缘计算的发展，利用摄像头或边缘设备对电力场景进行图像采集、缺陷识别和告警生成，已经成为智能运维的重要方向。国内外研究普遍关注 YOLO 等目标检测模型在缺陷识别中的应用，同时也越来越重视边云协同：边缘端负责近场采集和快速推理，云端负责数据汇聚、记录管理、可视化和后续处置。",
    )
    add_para(
        doc,
        "本项目不是单纯训练一个检测模型，而是把模型能力嵌入一个可运行的业务系统。对实际应用而言，模型输出需要经过接口封装、数据存储、告警关联、页面展示和人工复核，才能真正服务巡检业务。因此，系统架构、数据流和前后端工程实现是模型落地的关键环节。",
    )
    add_heading(doc, "1.2 项目总体介绍与分工", 2)
    add_para(
        doc,
        "项目总体目标是构建电力智能巡检工作台，实现边端设备在线上报、云端统一管理、监控员实时查看、巡检员现场补充记录、管理员查看系统概览的闭环流程。系统面向三类用户：监控员负责查看实时检测画面、推理结果和告警；巡检员负责上传现场检查报告并使用 Agent 辅助分析；管理员负责查看用户、节点、告警和报告概览，并可进入监控工作台。",
    )
    add_table(
        doc,
        ["角色/模块", "主要职责", "本人参与情况"],
        [
            ["系统架构", "确定边端、云端、前端和数据库之间的数据流与接口边界", "主要负责"],
            ["后端 app", "实现 REST API、数据库模型、业务服务、图片/报告保存、告警与记录处理", "主要负责"],
            ["前端 static", "实现监控员、巡检员、管理员页面以及实时数据展示", "主要负责"],
            ["边端推理", "采集图片、调用模型、上传证据和推理结果", "参与接口约定与联调"],
            ["模型训练", "训练和替换检测模型", "提供输出格式和部署对接约束"],
        ],
    )
    add_heading(doc, "1.3 本人负责工作概述", 2)
    add_bullets(
        doc,
        [
            "完成系统总体架构与数据流设计，形成顶层 DFD、0 层 DFD、关键 1 层 DFD 和数据字典。",
            "设计后端核心数据表和业务对象，包括用户、边缘节点、设备、检测数据、推理结果、巡检记录、告警事件、模型、命令、知识库和报告。",
            "实现 app 后端接口，覆盖登录、边端上报、命令下发、最新状态、巡检记录、告警处理、报告上传、Agent 对话、模型与知识库管理等功能。",
            "实现 static/dashboard 前端工作台，完成三类角色入口、实时图片展示、推理框绘制、告警表格、巡检记录分页、报告上传和管理员概览。",
            "修复多检测框展示、告警关联、设备 ID 传递和页面自动刷新等联调问题，使系统能够稳定展示边端每隔一段时间上报的新图片与推理结果。",
        ],
    )

    add_heading(doc, "2. 相关技术", 1)
    add_heading(doc, "2.1 边缘视觉推理与 YOLO 检测", 2)
    add_para(
        doc,
        "系统中的图像识别部分采用目标检测思路，边端模型对输入图片输出目标类别、置信度和边界框。为了降低模型替换对系统的影响，项目约定推理模块统一输出 list[dict]，每个检测结果包含 label、confidence、bbox、status、description 等字段。这样无论后续使用 YOLO、Faster R-CNN 还是分割模型，云端保存和前端展示逻辑都不需要大规模改动。",
    )
    add_code(
        doc,
        'infer(image_path) -> list[dict]\n\n{\n    "label": "insulator_defect",\n    "confidence": 0.83,\n    "bbox": [x, y, width, height],\n    "status": "warning",\n    "description": "可选说明"\n}',
    )
    add_heading(doc, "2.2 FastAPI 与 REST 接口", 2)
    add_para(
        doc,
        "云端后端使用 FastAPI 实现 REST API。边端通过 HTTP 主动上报心跳、证据图片、推理结果和原始检测数据；前端通过 API 获取最新状态、巡检记录、告警列表和管理员统计数据。REST 接口结构清晰，便于小组成员并行开发，也便于通过浏览器、脚本或边端程序进行联调。",
    )
    add_heading(doc, "2.3 SQLAlchemy 数据持久化", 2)
    add_para(
        doc,
        "后端使用 SQLAlchemy ORM 管理数据库表结构和对象关系，将 DFD 中的数据存储映射为实际表结构。巡检记录与推理结果、告警事件、设备和边缘节点之间建立关联，使页面能够从一条记录追溯到对应图片、检测框、异常状态和处理过程。",
    )
    add_heading(doc, "2.4 前端原生 Web 技术与实时轮询", 2)
    add_para(
        doc,
        "前端主要由 static/dashboard 下的 HTML、CSS 和 JavaScript 实现。页面通过 fetch 调用后端接口，并在 JavaScript 中绘制图片、检测框、表格和状态卡片。为了让监控员和管理员看到边端周期性上报的新图片，前端加入定时轮询逻辑，当前刷新间隔为 1 秒，并在浏览器控制台输出刷新日志，便于调试。",
    )
    add_heading(doc, "2.5 Agent 与 MCP 工具调用", 2)
    add_para(
        doc,
        "巡检员页面接入 Agent 对话能力，支持文本问答、图片理解以及系统数据查询。Agent 可以调用系统工具查询概览、巡检记录、告警详情和运行诊断；当未配置外部模型 API Key 时，后端提供本地 fallback 回复，保证页面功能不会因外部服务不可用而完全中断。",
    )

    add_heading(doc, "3. 系统设计与实现", 1)
    add_heading(doc, "3.1 总体架构设计", 2)
    add_para(
        doc,
        "系统采用三层架构：边缘节点负责采集图片并执行推理，FastAPI 云端负责接收数据、保存图片、生成巡检记录和告警，Web 工作台负责向不同角色展示数据并提供处理入口。整体链路如下：",
    )
    add_code(
        doc,
        "Atlas 200I DK / 摄像头\n"
        "  -> POST /api/edge/evidence 保存图片\n"
        "  -> POST /api/edge/inference 保存推理结果并生成告警/记录\n"
        "  -> FastAPI + SQLAlchemy 统一管理\n"
        "  -> static/dashboard 工作台实时展示与人工处理",
    )
    add_table(
        doc,
        ["DFD 数据存储", "系统实现", "说明"],
        [
            ["D1 用户与角色权限库", "users", "保存监控员、巡检员、管理员账号与角色"],
            ["D2 设备信息库", "edge_nodes, devices", "保存边缘节点、设备在线状态和最新上报时间"],
            ["D3 实时检测数据存储", "edge_detection_data, static/images", "保存原始检测数据和证据图片"],
            ["D4 边端推理结果库", "inference_results", "保存类别、置信度、bbox、推理耗时和模型信息"],
            ["D5 巡检记录库", "inspection_records", "形成可分页查看和处理的巡检记录"],
            ["D6 模型库", "models, model_deployments", "管理模型版本、框架和部署配置"],
            ["D7 告警与报表库", "alarm_events, reports", "保存告警事件和上传/生成的报告"],
            ["D8 设备资料/知识库", "inspection_knowledge", "为 Agent 和告警解释提供知识来源"],
        ],
    )
    add_heading(doc, "3.2 后端模块设计", 2)
    add_para(
        doc,
        "后端代码集中在 app 目录。controller.py 负责定义 API 路由并统一返回格式；schemas.py 使用 Pydantic 定义请求体；entitys.py 定义数据库 ORM 模型；services.py 承担核心业务逻辑；agent_service.py 负责 Agent 对话和工具调用。这样的分层使接口定义、数据校验、持久化对象和业务处理相互分离，便于多人协作开发。",
    )
    add_table(
        doc,
        ["文件", "主要内容"],
        [
            ["app/controller.py", "FastAPI 应用、静态资源挂载、登录、边端上报、记录、告警、报告、模型、知识库等接口"],
            ["app/services.py", "保存心跳、保存图片、保存推理结果、生成记录与告警、命令流转、报告上传、数据转字典"],
            ["app/entitys.py", "用户、节点、设备、推理结果、巡检记录、告警、模型、报告等 ORM 模型"],
            ["app/schemas.py", "接口输入数据结构，如 HeartbeatIn、InferenceIn、AgentChatIn、ReportIn 等"],
            ["app/agent_service.py", "Agent 消息构造、工具调用、图片转 data URL、fallback 回复"],
        ],
    )
    add_para(
        doc,
        "后端对边端开放的核心接口包括 /api/edge/heartbeat、/api/edge/evidence、/api/edge/inference、/api/edge/detection-data。对前端开放的核心接口包括 /api/status/latest、/api/records、/api/alerts/active、/api/reports/upload、/api/agent/chat、/api/edge/nodes 等。边端命令采用“云端创建命令、边端轮询领取、执行后 ACK”的方式实现，状态从 pending 流转到 running 和 success。",
    )
    add_heading(doc, "3.3 前端模块设计", 2)
    add_para(
        doc,
        "前端工作台集中在 static/dashboard。index.html 提供监控员、巡检员和管理员三类页面结构；styles.css 负责浅色网格风格、面板、表格和响应式布局；app.js 负责登录分流、接口请求、状态刷新、图片渲染、检测框绘制、表格渲染、报告上传、Agent 对话和管理员概览。",
    )
    add_table(
        doc,
        ["页面/视图", "实现内容"],
        [
            ["监控员工作台", "展示边端实时检测图片、推理图片、检测框、告警处理、边端记录和巡检员记录"],
            ["巡检员工作台", "支持 Report 上传、关联设备/记录、填写备注、Agent 对话和图片分析"],
            ["管理员控制台", "显示用户数、边缘节点数、活跃告警数、报告文件数，以及用户表和节点表"],
            ["实时刷新逻辑", "每 1 秒拉取最新状态、记录和告警；管理员页面同步刷新概览数据"],
        ],
    )
    add_code(
        doc,
        "const DASHBOARD_POLL_INTERVAL_MS = 1000;\n\n"
        "const [latest, records, alerts] = await Promise.all([\n"
        "  api('/api/status/latest'),\n"
        "  api(`/api/records?page=${recordsState.page}&page_size=${recordsState.pageSize}`),\n"
        "  api('/api/alerts/active'),\n"
        "]);",
    )
    add_heading(doc, "3.4 数据处理流程", 2)
    add_para(
        doc,
        "系统的一次典型巡检流程为：边端设备上报心跳，云端更新节点在线状态；边端采集图片后通过 evidence 接口上传，后端将图片保存到 static/images 并返回 image_uri；边端调用模型推理后通过 inference 接口上传检测结果；后端根据检测结果创建 InferenceResult、InspectionRecord 和 AlarmEvent；前端轮询 latest status、records 和 active alerts，将最新图片和检测框渲染到页面；监控员或巡检员处理告警后，后端更新告警状态和巡检记录状态。",
    )
    add_heading(doc, "3.5 关键实现细节", 2)
    add_bullets(
        doc,
        [
            "图片保存：后端对 node_id、record_id 和时间戳进行安全文件名处理，保存到 static/images，并通过 /images 静态路径给前端访问。",
            "多检测框展示：巡检记录不仅读取直接关联的第一条推理结果，还按 device_id、node_id、image_uri、infer_time 找回同一张图片上的全部推理结果。",
            "告警闭环：异常检测结果会生成告警事件，前端支持单个告警解除、按记录解除告警和完成记录处理。",
            "报告上传：巡检员上传文件后保存到 static/reports，并写入 reports 表；若关联巡检记录，则同步更新记录的文件 URI、文件名、处理状态和备注。",
            "管理员实时查看：管理员既能查看后台概览，也能切换到监控工作台；轮询逻辑覆盖 monitor 和 admin 两类页面。",
        ],
    )

    add_heading(doc, "4. 测试结果与分析", 1)
    add_heading(doc, "4.1 测试环境与测试方法", 2)
    add_para(
        doc,
        "测试主要在本地开发环境完成，后端通过 uvicorn 启动 FastAPI 服务，前端通过 /dashboard 访问。测试方式包括浏览器手动操作、接口联调脚本、页面 smoke test 和代码语法检查。重点验证边端上报数据能否被保存、最新状态能否被页面自动刷新、推理检测框能否完整展示、告警和记录是否能完成处理。",
    )
    add_table(
        doc,
        ["测试项", "测试内容", "结果"],
        [
            ["角色登录", "monitor、inspector、admin 三类账号登录后进入对应工作台", "通过"],
            ["边端图片保存", "调用 /api/edge/evidence 上传图片并在 static/images 生成文件", "通过"],
            ["推理结果展示", "调用 /api/edge/inference 上传 bbox，前端绘制检测框和标签", "通过"],
            ["自动刷新", "边端周期性上传后，页面无需手动刷新即可拉取最新状态", "通过，当前间隔 1 秒"],
            ["多 bbox 记录", "同一张图片多条检测结果在 latest status 与 records 中完整返回", "通过"],
            ["告警处理", "解除单个告警、按记录解除告警、完成巡检记录", "通过"],
            ["报告上传", "巡检员上传报告并关联设备/记录", "通过"],
            ["管理员概览", "管理员页面实时查看用户、节点、告警、报告统计", "通过"],
        ],
    )
    add_heading(doc, "4.2 运行结果截图", 2)
    if SCREENSHOT.exists():
        doc.add_picture(str(SCREENSHOT), width=Inches(6.2))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("图 4-1 监控员工作台运行测试截图：实时图片、检测框、推理结果和巡检记录展示")
        set_font(r, size=10)
    else:
        add_para(doc, "测试截图文件未找到，预期路径为：" + str(SCREENSHOT), first_line=False)
    add_heading(doc, "4.3 结果分析", 2)
    add_para(
        doc,
        "从测试结果看，系统已经实现从边端数据产生到云端保存、前端展示、人工处理的基本闭环。工作台能够把图片、bbox、状态、记录和告警组织到同一页面，监控员可以根据最新推理结果进行快速判断。巡检员上传报告后，系统能够把现场处理结果补充到同一条巡检记录中，避免推理结果和人工记录分散。管理员页面补充了系统级视角，便于查看当前用户、节点和告警规模。",
    )
    add_para(
        doc,
        "自动刷新逻辑解决了边端每隔约 10 秒上传一次图片而页面必须手动刷新的问题。当前刷新间隔设置为 1 秒，响应更及时，也方便调试；但在真实部署中需要根据设备数量和网络压力调整轮询频率，或者改为 WebSocket、SSE、MQTT 等事件驱动方式。",
    )

    add_heading(doc, "5. 实践过程中遇到的问题及解决方法", 1)
    add_table(
        doc,
        ["问题", "原因分析", "解决方法"],
        [
            [
                "前端必须手动刷新才能看到最新图片",
                "页面初始只在加载或点击刷新按钮时请求数据，没有持续拉取最新状态",
                "在 app.js 中加入定时轮询，周期调用 refreshDashboard；当前间隔为 1 秒，并在控制台输出 refresh 日志",
            ],
            [
                "管理员也需要实时查看，但最初只覆盖监控员页面",
                "轮询条件只判断 monitor，admin 后台概览没有进入自动刷新分支",
                "将可轮询角色扩展为 monitor 和 admin，管理员页调用 refreshAdminDashboard，切换到监控工作台后继续刷新",
            ],
            [
                "多个 bbox 不能完整展示",
                "一条巡检记录只直接关联第一条 InferenceResult，后续检测框没有被 record_to_dict 找回",
                "扩展 _inference_results_for_record，按 device_id、node_id、image_uri、infer_time 查回同一帧的全部推理结果",
            ],
            [
                "告警可能关联到错误推理结果",
                "批量保存检测结果和告警时，告警与具体检测项的一一对应关系不清晰",
                "保存推理结果时同步创建对应告警，保证 AlarmEvent 与当前检测结果关联",
            ],
            [
                "边端上传图片后前端看到的内容与预期不一致",
                "调试时容易混淆测试占位图、历史图片和最新 image_uri，同时浏览器缓存会增加判断难度",
                "统一以后端返回的 image_uri 作为前端图片来源，确认 static/images 保存文件，并通过刷新日志观察最新记录 ID 和图片路径",
            ],
            [
                "不同组员替换模型时容易影响云端和前端",
                "模型输出字段如果不统一，后端保存和前端绘制都需要跟着修改",
                "约定 infer(image_path) 的统一输出格式，模型内部可替换，但对系统暴露 label、confidence、bbox、status 等稳定字段",
            ],
        ],
    )

    add_heading(doc, "6. 总结与展望", 1)
    add_heading(doc, "6.1 工作总结", 2)
    add_para(
        doc,
        "本次实践中，本人主要完成了电力智能巡检系统的总体架构设计和前后端平台实现。架构层面，设计了边缘节点、云端后端、Web 工作台和数据库之间的数据流，明确了监控员、巡检员、管理员三类角色的功能边界。后端层面，实现了 app 目录中的 FastAPI 接口、SQLAlchemy 数据模型、图片与报告保存、推理结果入库、告警生成与处理、巡检记录管理、Agent 对话等功能。前端层面，实现了 static/dashboard 工作台，完成实时图片展示、推理框绘制、记录分页、告警处理、报告上传、管理员后台和 1 秒自动刷新。",
    )
    add_para(
        doc,
        "通过这些工作，系统从单独的模型推理脚本扩展为一个可协作、可联调、可演示的完整应用平台。组员可以在统一接口和数据结构下分别推进模型训练、边端采集、巡检业务和页面展示，降低了模块之间的耦合。",
    )
    add_heading(doc, "6.2 不足与优化方向", 2)
    add_bullets(
        doc,
        [
            "实时通信方面，当前主要依赖 HTTP 轮询，设备数量增加后会带来额外请求压力，后续可改为 WebSocket、SSE 或 MQTT。",
            "权限管理方面，目前角色分流和用户管理较基础，后续可加入更完整的 RBAC、操作审计和 Token 过期机制。",
            "模型管理方面，当前只完成基本版本和部署配置展示，后续可加入模型评估指标、灰度发布、回滚和边端部署状态追踪。",
            "数据存储方面，图片和报告当前保存在本地 static 目录，真实部署中可迁移到对象存储，并增加清理、归档和备份策略。",
            "告警策略方面，当前主要根据推理结果生成告警，后续可结合设备运行数据、历史记录和专家规则降低误报，提高严重缺陷的召回率。",
            "前端体验方面，可继续优化大屏布局、异常优先级排序、移动端巡检表单和离线场景下的数据暂存。",
        ],
    )

    add_heading(doc, "参考材料", 1)
    add_bullets(
        doc,
        [
            "20260618/dfd+dd.md：系统 DFD 顶层图、0 层图、关键 1 层图与数据字典。",
            "20260618/今日进展总结.md：前端工作台、后端接口、边云通信、Agent、报告上传和联调问题总结。",
            "intelligent-power-monitoring-system/app：后端 FastAPI、ORM 模型和业务服务实现。",
            "intelligent-power-monitoring-system/static/dashboard：前端三角色工作台实现。",
        ],
    )
    return doc


def main() -> None:
    TODAY.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
